# Define the standard proposal format
from datetime import datetime
from io import BytesIO
import os
import re
import traceback
from typing import Any, Dict
import uuid
from docx import Document
from flask import current_app
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
# from docx.oxml.ns import qn
from app.classes.ollama_service import OllamaService


DEFAULT_PROPOSAL_FORMAT = """
1. Background
   [ Brief overview of the terms of reference, Key benefits and outcomes, and scope summary. Must be at least 300 words]

2. Context
   [Detailed methodology, overview of the task. Must be at least 250 words. ]

3. Need of the Assignment
[explain the need of the assignment in one paragraph. addtionally, mention the role of agency in the assignment. The content must be at least 300 words long]

4. Work Package
[divide the assignment into 2-3 work packages. The work packages may further be divided into sub packages mentioned by headings as shown in the example, 
if required. Explanation for each sub packages paragrpah must be at least 200 words.]
   - work package 1
   - work package 2
   -...
   
[for example:
### Work Package 1: Data Collection and Analysis 
This work package involves gathering and analyzing data on both water demand and supply. It will be further divided into sub-packages:

    #### Demand Side Analysis:
        Focuses on collecting detailed information from industries, agricultural practices, livestock management, and domestic use patterns.Quantifies the total water demand using standardized metrics like LPCD for residential use, M/ha for agriculture, and GPM for industrial processes.

    #### Supply Side Assessment:
        Involves evaluating surface water availability from major rivers, lakes, and reservoirs; assessing groundwater levels through well data and aquifer studies; and measuring rainwater runoff using catchment area analysis. Utilizes hydrological models to predict supply based on climate patterns and long-term forecasts.
]

5. Deliverables
[mention the deliverables expected as output from the agency]

6. Timeline
[for example: The duration of the contract will be from 01 Mar 2025 to 30 Apr 2025]

7. Expert Requirements
[List the expected experts along with their respective educational qualfications and experience required]

7.Person days
[example: The assignment will require at lease 40 person days ]
"""

class Orchestrator():
    def __init__(self):
        """Initialize the orchestrator with OllamaService."""
        self.ollama_service = OllamaService()
        self.output_dir = current_app.config["UPLOAD_FOLDER"]
        self.proposal_format = DEFAULT_PROPOSAL_FORMAT
        
    def generate_proposal(self, task_id, context: str, progress_queues,  model="gemma3:12b", stream=False) -> Dict[str, Any]:
        """Orchestrate the entire proposal generation process."""
        # return self.ollama_service.generate_draft_tor(context, self.proposal_format, model=model, stream=stream)
        # task_id = str(uuid.uuid4())
        
        try:
        #     # Step 1: Creator generates draft proposal
        #     # logger.info(f"[{task_id}] Starting proposal generation process")
            progress_queues[task_id].put({"progress": 25, "status": "Upload complete. Creating Draft..."})
            draft_proposal = self.ollama_service.create_terms_of_reference(context, self.proposal_format)
            
            # Step 2: Reviewer reviews the draft
            # logger.info(f"[{task_id}] Draft created, starting review")
            progress_queues[task_id].put({"progress": 50, "status": "Draft complete. Reviewing Draft..."})
            review_feedback = self.ollama_service.review_terms_of_reference(draft_proposal, context, self.proposal_format)
            
            # Step 3: Creator revises based on feedback
            # logger.info(f"[{task_id}] Review complete, starting revision")
            progress_queues[task_id].put({"progress": 75, "status": "Reviewing complete. Revising document..."})
            final_proposal = self.ollama_service.revise_terms_of_reference(draft_proposal, review_feedback, context, self.proposal_format)
            
            # Step 4: Evaluator evaluates the final proposal
            # logger.info(f"[{task_id}] Revision complete, starting evaluation")
            progress_queues[task_id].put({"progress": 99, "status": "Revision complete. Evaluating document..."})
            evaluation = self.ollama_service.evaluate_terms_of_reference(final_proposal, context, self.proposal_format)
            
            # Parse evaluation score
            try:
                score_text = re.search(r'SCORE:\s*(\d+)', evaluation)
                score = int(score_text.group(1)) if score_text else 3
                # logger.info(f"[{task_id}] Parsed score: {score}")
            except Exception as e:
                # logger.warning(f"[{task_id}] Could not parse evaluation score: {str(e)}")
                score = 3
            
            # Create output document
            filename, filestream = self.create_docx(final_proposal, context, review_feedback, evaluation, score, task_id)
            
            return {
                "task_id": task_id,
                "draft_proposal": draft_proposal,
                "review_feedback": review_feedback,
                "final_proposal": final_proposal,
                "evaluation": evaluation,
                "score": score,
                "filename": filename,
                "file": filestream,
                "status": "complete"
            }
            
        except Exception as e:
            error_details = traceback.format_exc()
            # logger.error(f"[{task_id}] Error in proposal generation: {error_details}")
            return {
                "task_id": task_id,
                "error": str(e),
                "status": "failed",
                "details": error_details if current_app.config["DEBUG"] else None
            }
    
    def create_docx(self, proposal: str, context: str, feedback: str, evaluation: str, 
                score: int, task_id: str) -> str:
        """Create a professional DOCX file with the final proposal, converting markdown to docx formatting."""

        doc = Document()
        
        # Add document properties
        doc.core_properties.author = "JiM"
        doc.core_properties.title = "Terms of Reference"
        
        # Define document styles
        styles = doc.styles
        
        # Modify the Normal style
        style_normal = styles['Normal']
        font = style_normal.font # type: ignore
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Create section heading styles
        for i, heading_name in enumerate(['Title', 'Section Heading', 'Subsection Heading']):
            style_name = f'Custom {heading_name}'
            if style_name not in styles:
                style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = styles['Heading ' + str(i+1)] # type: ignore
                style.font.name = 'Calibri' # type: ignore
                style.font.size = Pt(16 - i*2) # type: ignore
                style.font.bold = True # type: ignore
                style.font.color.rgb = RGBColor(0, 70, 127)  # type: ignore # Blue color
        
        # Create additional markdown styles
        if 'Code Block' not in styles:
            code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Courier New' # type: ignore
            code_style.font.size = Pt(10) # type: ignore
            code_style.paragraph_format.left_indent = Inches(0.5) # type: ignore
            code_style.paragraph_format.right_indent = Inches(0.5) # type: ignore
            code_style.paragraph_format.space_before = Pt(6) # type: ignore
            code_style.paragraph_format.space_after = Pt(6) # type: ignore
        
        if 'List Item' not in styles:
            list_style = styles.add_style('List Item', WD_STYLE_TYPE.PARAGRAPH)
            list_style.base_style = styles['Normal'] # type: ignore
            list_style.paragraph_format.left_indent = Inches(0.25) # type: ignore
            list_style.paragraph_format.first_line_indent = Inches(-0.25) # type: ignore
        
        # Add title page
        title = doc.add_paragraph("Terms of Reference", style='Custom Title')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = datetime.now().strftime("%B %d, %Y")
        date_paragraph = doc.add_paragraph(f"Generated: {timestamp}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add quality score with star rating
        score_paragraph = doc.add_paragraph(f"Quality Score: {score}/5 ")
        score_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add stars for visual score representation
        run = score_paragraph.add_run("★" * score + "☆" * (5 - score))
        run.font.color.rgb = RGBColor(255, 191, 0)  # Gold color
        
        # Add a page break
        # doc.add_page_break()
        
        # Add table of contents header
        doc.add_paragraph("Table of Contents", style='Custom Section Heading')
        
        # Extract section headings from the proposal for TOC
        toc_items = []
        for line in proposal.split('\n'):
            # Match markdown headings (# Heading) or numbered sections (1. Section)
            if re.match(r'^#+\s+', line.strip()) or re.match(r'^\d+\.', line.strip()):
                # Clean up markdown formatting for TOC
                clean_line = re.sub(r'^#+\s+', '', line.strip())
                toc_items.append(clean_line)
        
        # Create TOC as a list
        for item in toc_items:
            p = doc.add_paragraph()
            p.add_run(item)
            p.paragraph_format.left_indent = Inches(0.25)
        
        # Add a page break after TOC
        # doc.add_page_break()
        
        # Add the proposal content - parse markdown for better formatting
        doc.add_paragraph("Proposal Content", style='Custom Section Heading')
        
        # Parse and format markdown content
        lines = proposal.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
                
            # Process markdown headings (e.g., # Heading)
            heading_match = re.match(r'^(#+)\s+(.*)', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                
                if level == 1:
                    p = doc.add_paragraph(text, style='Custom Section Heading')
                else:
                    p = doc.add_paragraph(text, style='Custom Subsection Heading')
                    
                i += 1
                continue
            
            # Process numbered headings (e.g., 1. Heading)
            if re.match(r'^\d+\.', line):
                p = doc.add_paragraph(line, style='Custom Section Heading')
                i += 1
                continue
                
            # Process lists
            list_match = re.match(r'^(\s*)([\*\-\+]|\d+\.)\s+(.*)', line)
            if list_match:
                indent = len(list_match.group(1))
                marker = list_match.group(2)
                text = list_match.group(3)
                
                p = doc.add_paragraph(style='List Item')
                p.add_run(f"{marker} {text}")
                p.paragraph_format.left_indent = Inches(0.25 + (indent * 0.25))
                
                i += 1
                continue
                
            # Process code blocks
            if line.startswith('```'):
                code_lang = line[3:]
                code_content = []
                i += 1
                
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_content.append(lines[i])
                    i += 1
                    
                if i < len(lines):  # Skip the closing ```
                    i += 1
                    
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph(code_text, style='Code Block')
                continue
                
            # Process horizontal rules
            if re.match(r'^(\*\*\*|\-\-\-|___)', line):
                doc.add_paragraph().add_run().add_break()
                p = doc.add_paragraph()
                p.paragraph_format.border_bottom = True # type: ignore
                i += 1
                continue
                
            # Process regular paragraph with inline formatting
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.25)
            
            # Handle inline formatting
            j = 0
            text = line
            
            # Process bold, italics, code, etc.
            while j < len(text):
                # Bold: **text** or __text__
                bold_match = re.match(r'(\*\*|__)(.*?)(\*\*|__)', text[j:])
                if bold_match:
                    run = p.add_run(bold_match.group(2))
                    run.bold = True
                    j += len(bold_match.group(0))
                    continue
                    
                # Italic: *text* or _text_
                italic_match = re.match(r'(\*|_)(.*?)(\*|_)', text[j:])
                if italic_match:
                    run = p.add_run(italic_match.group(2))
                    run.italic = True
                    j += len(italic_match.group(0))
                    continue
                    
                # Inline code: `code`
                code_match = re.match(r'`(.*?)`', text[j:])
                if code_match:
                    run = p.add_run(code_match.group(1))
                    run.font.name = 'Courier New'
                    j += len(code_match.group(0))
                    continue
                    
                # Regular text (no formatting)
                if j < len(text):
                    # Find the next formatting marker
                    next_marker = float('inf')
                    for marker in ['**', '__', '*', '_', '`']:
                        pos = text.find(marker, j) # type: ignore
                        if pos != -1 and pos < next_marker:
                            next_marker = pos
                    
                    if next_marker == float('inf'):
                        # No more formatting, add the rest of the text
                        p.add_run(text[j:])
                        break
                    else:
                        # Add text up to the next formatting marker
                        p.add_run(text[j:next_marker])
                        j = next_marker
            
            i += 1
        
        # Add appendices
        doc.add_page_break()
        doc.add_paragraph("Appendices", style='Custom Section Heading')
        
        # Appendix A: Original Context
        doc.add_paragraph("Appendix A: Original Context", style='Custom Subsection Heading')
        # Parse markdown in context
        self._add_markdown_content(doc, context)
        
        # Appendix B: Review Feedback
        doc.add_paragraph("Appendix B: Review Feedback", style='Custom Subsection Heading')
        # Parse markdown in feedback
        self._add_markdown_content(doc, feedback)
        
        # Appendix C: Evaluation
        doc.add_paragraph("Appendix C: Evaluation", style='Custom Subsection Heading')
        # Parse markdown in evaluation
        self._add_markdown_content(doc, evaluation)
        
        # Generate unique filename
        filename = f"tor_{task_id}.docx"
        filepath = os.path.join(self.output_dir, filename)
        file_stream = BytesIO()
        
        # Save document
        doc.save(filepath)
        doc.save(file_stream)
        # logger.info(f"Created proposal document: {filepath}")
        
        return filename, file_stream # type: ignore

    def _add_markdown_content(self, doc, markdown_text):
        """Helper method to parse and add markdown content to the document."""
        if not markdown_text:
            doc.add_paragraph("No content provided.")
            return
            
        # Simple markdown parsing for appendices
        for line in markdown_text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Check if heading
            if line.startswith('####'):
                level = len(re.match(r'^(####+)', line).group(1)) # type: ignore
                text = line[level+1:]
                if level == 1:
                    doc.add_paragraph(text, style='Custom Section Heading')
                else:
                    doc.add_paragraph(text, style='Custom Subsection Heading')
            elif line.startswith('#'):
                level = len(re.match(r'^(#+)', line).group(1)) # type: ignore
                text = line[level+1:]
                if level == 1:
                    doc.add_paragraph(text, style='Custom Section Heading')
                else:
                    doc.add_paragraph(text, style='Custom Subsection Heading')
            # # Check if bold item
            # elif re.match(r'^[\*\*\-\+]|\d+\.', line):
            #     # Remove the ** pair using regex
            #     cleaned_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            #     p = doc.add_paragraph(style='Custom Subsection Heading')
            #     p.add_run(cleaned_line)
             # Check if list text
            elif re.match(r'^[\*\-\+]|\d+\.', line):
                p = doc.add_paragraph(style='List Item')
                p.add_run(line)
            # Regular paragraph
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.first_line_indent = Inches(0.25)
    
    # def create_docx(self, proposal: str, context: str, feedback: str, evaluation: str, 
    #                 score: int, task_id: str) -> str:
    #     """Create a professional DOCX file with the final proposal."""
    #     doc = Document()
        
    #     # Add document properties
    #     doc.core_properties.author = "JiM"
    #     doc.core_properties.title = "Terms of Reference"
        
    #     # Define document styles
    #     styles = doc.styles
        
    #     # Modify the Normal style
    #     style_normal = styles['Normal']
    #     font = style_normal.font
    #     font.name = 'Calibri'
    #     font.size = Pt(11)
        
    #     # Create section heading styles
    #     for i, heading_name in enumerate(['Title', 'Section Heading', 'Subsection Heading']):
    #         style_name = f'Custom {heading_name}'
    #         if style_name not in styles:
    #             style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    #             style.base_style = styles['Heading ' + str(i+1)]
    #             style.font.name = 'Calibri'
    #             style.font.size = Pt(16 - i*2)
    #             style.font.bold = True
    #             style.font.color.rgb = RGBColor(0, 70, 127)  # Blue color
        
    #     # Add title page
    #     title = doc.add_paragraph("Terms of Reference", style='Custom Title')
    #     title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    #     # Add timestamp
    #     timestamp = datetime.now().strftime("%B %d, %Y")
    #     date_paragraph = doc.add_paragraph(f"Generated: {timestamp}")
    #     date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    #     # Add quality score with star rating
    #     score_paragraph = doc.add_paragraph(f"Quality Score: {score}/5 ")
    #     score_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    #     # Add stars for visual score representation
    #     run = score_paragraph.add_run("★" * score + "☆" * (5 - score))
    #     run.font.color.rgb = RGBColor(255, 191, 0)  # Gold color
        
    #     # Add a page break
    #     doc.add_page_break()
        
    #     # Add table of contents header
    #     doc.add_paragraph("Table of Contents", style='Custom Section Heading')
        
    #     # Extract section headings from the proposal format to create TOC
    #     toc_items = []
    #     for line in proposal.split('\n'):
    #         if re.match(r'^\d+\.', line.strip()):
    #             toc_items.append(line.strip())
        
    #     # Create TOC as a list
    #     for item in toc_items:
    #         p = doc.add_paragraph()
    #         p.add_run(item)
    #         p.paragraph_format.left_indent = Inches(0.25)
        
    #     # Add a page break after TOC
    #     doc.add_page_break()
        
    #     # Add the proposal content - parse sections for better formatting
    #     doc.add_paragraph("Proposal Content", style='Custom Section Heading')
        
    #     # Parse and format sections
    #     current_section = None
    #     for line in proposal.split('\n'):
    #         line = line.strip()
    #         if not line:
    #             continue
                
    #         # Check if this is a main section heading
    #         if re.match(r'^\d+\.', line):
    #             current_section = doc.add_paragraph(line, style='Custom Section Heading')
    #         # Check if this is a subsection heading
    #         elif re.match(r'^\s*-', line) and current_section:
    #             doc.add_paragraph(line, style='Custom Subsection Heading')
    #         # Otherwise it's regular content
    #         else:
    #             p = doc.add_paragraph(line)
    #             p.paragraph_format.first_line_indent = Inches(0.25)
        
    #     # Add appendices
    #     doc.add_page_break()
    #     doc.add_paragraph("Appendices", style='Custom Section Heading')
        
    #     # Appendix A: Original Context
    #     doc.add_paragraph("Appendix A: Original Context", style='Custom Subsection Heading')
    #     doc.add_paragraph(context)
        
    #     # Appendix B: Review Feedback
    #     doc.add_paragraph("Appendix B: Review Feedback", style='Custom Subsection Heading')
    #     doc.add_paragraph(feedback)
        
    #     # Appendix C: Evaluation
    #     doc.add_paragraph("Appendix C: Evaluation", style='Custom Subsection Heading')
    #     doc.add_paragraph(evaluation)
        
    #     # Generate unique filename
    #     filename = f"tor_{task_id}.docx"
    #     filepath = os.path.join(self.output_dir, filename)
    #     file_stream = BytesIO()
        
    #     # Save document
    #     doc.save(filepath)
    #     doc.save(file_stream)
    #     # logger.info(f"Created proposal document: {filepath}")
        
    #     return filename, file_stream