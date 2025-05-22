import json
from typing import Dict, List

from docx import Document
import ollama
from app.classes.ollama_service import OllamaService


class ChatStudio():
    def __init__(self, model:str = 'llama3.latest', style:str = 'normal', file_path:str=''):
        self.model = model
        self.style = style.lower()
        self.file_path = file_path
        self.temperature = 0.7
        self.system_prompt = self.get_system_prompt()
    
    def _chat_completion(self, messages: List[Dict[str, str]], stream=False):
        """Generate a response using Ollama chat API with caching."""
        # Create a deterministic representation of messages for caching
        ollama_service = OllamaService()
        messages_str = json.dumps(messages, sort_keys=True)
        cache_key = ollama_service._cache_key(messages_str, "")
        cached_response = ollama_service._check_cache(cache_key)
        if cached_response:
            return cached_response
            
        try:
            return ollama.chat(
                model=self.model,
                messages=messages,
                options={"temperature": self.temperature},
                stream=stream,
            )
            # result = response['message']['content']
            # ollama_service._save_to_cache(cache_key, result)
            # return result
        
        except Exception as e:
            error_msg = f"Error in chat completion: {str(e)}"
            return error_msg
        
    def _text_generation(self, prompt:str, system_prompt:str, stream=False):
        """Generate a response using Ollama text generation API with caching."""
        # Create a deterministic representation of the prompt for caching
        ollama_service = OllamaService()
        cache_key = ollama_service._cache_key(prompt, system_prompt)
        cached_response = ollama_service._check_cache(cache_key)
        if cached_response:
            return cached_response
            
        try:
            return ollama.generate(
                model=self.model,
                prompt=prompt,
                system=self.get_system_prompt(),
                options={"temperature": self.temperature},
                stream=stream
            )
            # result = response['response']
            # ollama_service._save_to_cache(cache_key, result)
            # return result
        except Exception as e:
            error_msg = f"Error generating response: {str(e)}"
            return error_msg
        
    def get_bot_response(self, message, file_path=None): 
        """
        Get the system prompt based on the style
        """ 
        if file_path:
            file_type = file_path.split('.')[-1]
            file_content = ''
            if file_type == 'pdf':
                file_content = self.pdf_to_text(file_path)
            elif file_type == 'docx':
                file_content = self.docx_to_text(file_path)
            elif file_type == 'txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    file_content = file.read()
            else:
                raise ValueError("Unsupported file type") 
            
            prompt = self.get_file_query_prompt(file_content, file_type, message)
            return self._text_generation(prompt = prompt, system_prompt=self.get_system_prompt(), stream=True)
        
        else:
            return self._chat_completion(
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": message}
                ],
                stream=True
            )
        
    def get_system_prompt(self):
        """
            Get the system prompt based on the style
        """
        if self.style == 'normal':
            return """
            Respond in a balanced, helpful, and natural tone. Use conversational language while
            providing informative and comprehensive answers. Balance technical depth with
            accessibility, use examples when helpful, and structure longer responses for readability.
            Keep your tone friendly and engaging without being overly casual or formal.
            """
        elif self.style == 'formal':
            return """
            Respond in a professional, academic tone. Use precise language and maintain a respectful,
            authoritative style throughout. Avoid contractions, colloquialisms, and casual expressions.
            Structure your responses carefully with clear organization. Use proper terminology and
            maintain a scholarly approach. Be thorough in your explanations while remaining dignified
            and measured in your expression.
            """
        elif self.style == 'explanatory':
            return """
            Focus on detailed explanations with educational value. Break down complex concepts into
            understandable components. Use analogies, examples, and step-by-step explanations to
            illustrate points. Define technical terms when they first appear. Structure responses with
            clear headings and progressive disclosure of information. Connect new information to
            familiar concepts and emphasize the 'why' behind facts and processes.
            """
        elif self.style == 'concise':
            return """
            Provide brief, direct responses that deliver essential information efficiently. Prioritize
            brevity while maintaining clarity and completeness. Use short sentences and paragraphs.
            Avoid unnecessary details, examples, or preambles. Get straight to the point and focus on
            key facts and actionable information. Use bullet points for lists rather than lengthy
            explanations.
            """
        else:
            # Default to normal if unknown style is provided
            return """
            Respond in a balanced, helpful, and natural tone. Use conversational language while
            providing informative and comprehensive answers. Balance technical depth with
            accessibility, use examples when helpful, and structure longer responses for readability.
            Keep your tone friendly and engaging without being overly casual or formal.
            """
        
        
    def get_file_query_prompt(self, file_content, file_type, user_query):
        """
        Create a prompt for responding to a user query about an uploaded file
        
        Parameters:
        - file_content: The extracted text content from the file
        - file_type: The type of file (pdf, docx, txt)
        - user_query: The user's question about the file
        
        Returns: A prompt string for the AI to process
        """
        
        # Create context about the document
        file_type_descriptions = {
            'pdf': "a PDF document",
            'docx': "a Word document",
            'txt': "a text file"
        }
        
        file_description = file_type_descriptions.get(file_type.lower(), "a document")
        
        prompt = f"""
        The user has uploaded {file_description} and has a question about it. 
        
        The content of the document is enclosed between triple backticks:
        ```
        {file_content}
        ```
        
        The user's question is:
        "{user_query}"
        
        Please provide a helpful, accurate response to their question based on the document's content.
        If the document doesn't contain information relevant to their query, politely explain this.
        If they've asked for a summary, provide a concise overview of the key points.
        If they've asked about specific information, cite the relevant sections from the document.
        If the document is very technical, explain concepts in clear, accessible language.
        
        Base your response solely on the content of the document and your general knowledge.
        """
        
        return prompt
    
    def docx_to_text(self, file_path):
        """
        Convert a document to text using the appropriate method based on the file type
        """
        doc = Document(file_path)
        # Extract text from paragraphs
        paragraphs_text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                tables_text.append(' | '.join(row_text))
        
        # Combine all text
        file_content = '\n\n'.join(paragraphs_text + tables_text)
        return file_content
    
    def pdf_to_text(self, file_path):
        """
        Convert a PDF document to text using PyPDF2
        """
        raise ValueError("Unsupported file type")
        