from docx import Document
from docx.oxml.ns import qn
from app.classes.ollama_service import OllamaService


class DocumentOperations():
    def __init__(self, paragraphs=[], prompt=""):
        self.paragraphs = paragraphs
        self.prompt = prompt
        if self.prompt == "":
            self.prompt = '''
            You are a proofreading expert. Please reivew the text for grammar, spelling, clairty, consistency, and technical accuracy. 
            Provide only the corrected version of the text without any additional comments, explainations, or formatting changes 
            beyond what is necessary for correctness.
            '''
        
    def read_docx(self, filepath, progress_queues, task_id):
        document = Document(filepath)
        processed_document = Document()
        self.paragraphs = []
        for index,element in enumerate(document.element.body):
            element_count = len(document.element.body)
            if element.tag == qn('w:p'):
                para = element.text
                word_count = len(para.split(" "))
                if word_count >= 1:
                    if  word_count <= 5:
                        if para == '':
                            continue
                        else:
                            # self.paragraphs.append(para)
                            processed_document.add_heading(para)
                    else:
                        processed_para = self.proofread_content(para)
                        bot_response = processed_para['response'] # type: ignore
                        processed_document.add_paragraph(bot_response)
                        progress_percent = int((index + 1) / element_count * 100)
                        progress_queues[task_id].put({"progress": progress_percent, "status": "Reading document..."})
                        # self.paragraphs.append(bot_response)
            elif element.tag == qn('w:tbl'):
                pass # its a table
            elif element.tag == qn('w:drawing'):
                pass # it is an image
            
        
        return processed_document   
    
    def proofread_content(self, content):
        ollama_service = OllamaService()
        return ollama_service.generate(prompt = self.prompt + content, stream = False)
    
    def save_docx(self, filepath, paragraphs):
        document = Document()
        for para in paragraphs:
            if para == '\n':
                continue
            else:
                word_count = len(para.split(" "))
                if word_count >= 1:
                    if word_count <= 5:
                        document.add_heading(para)
                    else:
                        document.add_paragraph(para)
        
        document.save(filepath)